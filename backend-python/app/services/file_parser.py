"""
文件解析服务
支持 PDF、Word(.docx)、TXT、Markdown 文件解析
"""
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def parse_file(file_path: str, file_type: str) -> str:
    """根据文件类型调用对应的解析器"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    file_type = file_type.lower()

    if file_type == '.pdf':
        return parse_pdf(file_path)
    elif file_type == '.docx':
        return parse_docx(file_path)
    elif file_type in ('.txt', '.md', '.markdown'):
        return parse_text(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")


def parse_pdf(file_path: str) -> str:
    """解析 PDF 文件"""
    try:
        import fitz  # PyMuPDF

        text_parts = []
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                text = page.get_text().strip()
                if text:
                    text_parts.append(text)

        result = '\n\n'.join(text_parts)
        if not result:
            raise ValueError("PDF 文件内容为空或无法提取文本")

        logger.info(f"PDF 解析完成: {file_path}, 共 {len(text_parts)} 页")
        return result

    except ImportError:
        raise ImportError("PDF 解析依赖未安装: pip install PyMuPDF")


def parse_docx(file_path: str) -> str:
    """解析 Word 文件"""
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        result = '\n'.join(text_parts)
        if not result:
            raise ValueError("Word 文件内容为空")

        logger.info(f"Word 解析完成: {file_path}")
        return result

    except ImportError:
        raise ImportError("Word 解析依赖未安装: pip install python-docx")


def parse_text(file_path: str) -> str:
    """解析纯文本文件（TXT/Markdown）"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            if content.strip():
                logger.info(f"文本解析完成: {file_path}, 编码: {encoding}")
                return content
        except (UnicodeDecodeError, UnicodeError):
            continue

    raise ValueError(f"无法读取文件，尝试了以下编码: {', '.join(encodings)}")
